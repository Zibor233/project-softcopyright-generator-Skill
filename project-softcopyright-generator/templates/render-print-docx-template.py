#!/usr/bin/env python3
"""
打印版 DOCX 转换模板

用途：
- 将草稿目录中的 markdown 内容转换为正式资料目录中的可打印 docx
- 适用于正式资料生成阶段的系统说明文档和源代码文档排版

主要配置字段：
- system_name
- workdir
- draft_dir_name
- final_dir_name
- system_md_path
- source_md_path
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PRINT_DOCX_CONFIG = {
    "system_name": "{system_name}",
    "workdir": "D:/your-project/软著资料",
    "draft_dir_name": "草稿",
    "final_dir_name": "正式资料",
    "system_md_path": None,
    "source_md_path": None,
}


def set_run_font(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.bold = bold


def setup_page(document: Document, narrow: bool = False):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if narrow:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
    else:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_title(document: Document, text: str, size: float = 18):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Microsoft YaHei", size, True)


def add_heading(document: Document, text: str, level: int):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    size = 16 if level == 1 else 14 if level == 2 else 12
    run = p.add_run(text)
    set_run_font(run, "黑体", "Microsoft YaHei", size, True)


def add_normal_paragraph(document: Document, text: str, centered: bool = False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 12)


def add_code_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, "等线", "Courier New", 8)


def resolve_paths(config):
    workdir = Path(config["workdir"])
    draft_dir = workdir / config["draft_dir_name"]
    final_dir = workdir / config["final_dir_name"]
    return {
        "system_md_path": Path(config["system_md_path"]) if config["system_md_path"] else draft_dir / f"{config['system_name']}_系统说明文档.md",
        "source_md_path": Path(config["source_md_path"]) if config["source_md_path"] else draft_dir / f"{config['system_name']}_源代码文档.md",
        "system_docx_path": final_dir / f"{config['system_name']}_系统说明文档.docx",
        "source_docx_path": final_dir / f"{config['system_name']}_源代码文档.docx",
    }


def build_system_docx(src: Path, dst: Path):
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_page(doc, narrow=False)

    for raw in lines:
        line = raw.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            add_title(doc, line[2:].strip(), 18)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif line == "【图片预留截图位置】":
            add_normal_paragraph(doc, line, centered=True)
        else:
            add_normal_paragraph(doc, line)

    doc.save(dst)
    return dst


def build_source_docx(src: Path, dst: Path):
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_page(doc, narrow=True)

    in_code = False
    first_page = True

    for raw in lines:
        line = raw.rstrip("\n")
        if line.startswith("# "):
            add_title(doc, line[2:].strip(), 18)
            continue
        if line.startswith("### 文件：") or line.startswith("文件："):
            add_heading(doc, line.replace("### ", "", 1).strip(), 3)
            continue
        if line.startswith("#### 第") or line.startswith("第 ") and line.endswith(" 页"):
            if not first_page:
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
            add_normal_paragraph(doc, line.replace("#### ", "", 1).strip(), centered=True)
            first_page = False
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            if in_code:
                add_code_paragraph(doc, "")
            continue
        if in_code:
            add_code_paragraph(doc, line)

    doc.save(dst)
    return dst


def render_print_docx(config=None):
    settings = config or PRINT_DOCX_CONFIG
    paths = resolve_paths(settings)
    outputs = []

    if paths["system_md_path"].exists():
        outputs.append(build_system_docx(paths["system_md_path"], paths["system_docx_path"]))
    if paths["source_md_path"].exists():
        outputs.append(build_source_docx(paths["source_md_path"], paths["source_docx_path"]))

    for item in outputs:
        print(f"✅ 已生成打印版文档: {item}")
    return outputs


if __name__ == "__main__":
    render_print_docx()
